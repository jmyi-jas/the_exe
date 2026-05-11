# -*- coding: utf-8 -*-
"""
法院工具箱
功能1：法院文书下载
功能2：起诉状信息提取（含扫描版 PDF OCR 支持，EasyOCR 版）
功能3：初筛 - 判决/裁定书文件夹重命名
"""

import os
import re
import sys
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from collections import defaultdict
from pathlib import Path
from typing import Optional, Tuple          # ← 新增此行

import pandas as pd
import requests
from urllib.parse import urlparse, parse_qs

# ══════════════════════════════════════════════
# 公共依赖检测
# ══════════════════════════════════════════════
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import easyocr
    HAS_EASYOCR = True
except ImportError:
    HAS_EASYOCR = False

try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False

_ocr_reader_instance = None


def _get_ocr_reader(log=print):
    global _ocr_reader_instance
    if _ocr_reader_instance is None:
        log("    [OCR] 首次初始化 EasyOCR 模型，请稍候（约 10~30 秒）...")
        _ocr_reader_instance = easyocr.Reader(
            ['ch_sim', 'en'], gpu=False, verbose=False
        )
        log("    [OCR] 模型加载完成")
    return _ocr_reader_instance


# ══════════════════════════════════════════════
# 一、文书下载 - 业务逻辑
# ══════════════════════════════════════════════

API_URL = "https://zxfw.court.gov.cn/yzw/yzw-zxfw-sdfw/api/v1/sdfw/getWsListBySdbhNew"
session = requests.Session()
session.trust_env = False

def dl_extract_case_number(text):
    # 兼容中文括号（）和英文括号()
    pattern = r'[（(]\d{4}[）)][^（(]*?\d+号'
    match = re.search(pattern, text)
    return match.group(0).strip() if match else None

def dl_sanitize(name):
    name = name.replace(':', '.')
    return re.sub(r'[\\/*?:"<>|]', '_', name)


def dl_extract_params(sms_text):
    match = re.search(r'https?://[^\s]+', sms_text)
    if not match:
        return None
    try:
        parsed = urlparse(match.group(0))
        if parsed.fragment and '?' in parsed.fragment:
            qs     = parsed.fragment.split('?', 1)[1]
            params = __import__('urllib.parse', fromlist=['parse_qs']).parse_qs(qs)
            qdbh   = params.get('qdbh',  [None])[0]
            sdbh   = params.get('sdbh',  [None])[0]
            sdsin  = params.get('sdsin', [None])[0]
            if qdbh and sdbh and sdsin:
                return {'qdbh': qdbh, 'sdbh': sdbh, 'sdsin': sdsin}
    except Exception:
        pass
    return None


def dl_fetch_files(params):
    headers = {
        'Accept': '*/*',
        'Content-Type': 'application/json',
        'Origin': 'https://download-tool.boyifuture.com',
        'Referer': 'https://download-tool.boyifuture.com/',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)',
    }
    try:
        resp = session.post(API_URL, json=params, headers=headers, timeout=10)
        if resp.status_code != 200:
            return []
        data = resp.json()
        if data.get('code') == 200 and isinstance(data.get('data'), list):
            return data['data']
    except Exception:
        pass
    return []


def dl_download_file(url, save_path):
    try:
        r = requests.get(url,
                         headers={'User-Agent': 'Mozilla/5.0'},
                         stream=True, timeout=30)
        r.raise_for_status()
        with open(save_path, 'wb') as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
        return True
    except Exception:
        return False


def dl_process_sms(sms_text, folder_path, log):
    try:
        with open(os.path.join(folder_path, "短信内容.txt"), 'w', encoding='utf-8') as f:
            f.write(sms_text)
    except Exception as e:
        log(f"  保存短信失败: {e}")

    params = dl_extract_params(sms_text)
    if not params:
        log("  未找到下载参数，跳过")
        return

    files = dl_fetch_files(params)
    if not files:
        log("  未获取到文件列表")
        return

    log(f"  发现 {len(files)} 个文件")
    for idx, fi in enumerate(files, 1):
        name      = fi.get('c_wsmc', f'文书_{idx}')
        ext       = fi.get('c_wjgs') or 'pdf'
        safe_name = dl_sanitize(f"{name}.{ext}")
        url       = fi.get('wjlj')
        if not url:
            log(f"    {safe_name} 无链接，跳过")
            continue
        log(f"    下载: {safe_name}")
        ok = dl_download_file(url, os.path.join(folder_path, safe_name))
        log(f"    {'✓ 成功' if ok else '✗ 失败'}")


def dl_run_task(excel_path, content_col, time_col, output_dir, log, done):
    try:
        df = pd.read_excel(excel_path, sheet_name=0, dtype=str)
    except Exception as e:
        log(f"❌ 读取Excel失败：{e}")
        done()
        return

    for col in [content_col, time_col]:
        if col not in df.columns:
            log(f"❌ 找不到列：{col}")
            log(f"   当前列名：{list(df.columns)}")
            done()
            return

    os.makedirs(output_dir, exist_ok=True)
    created = set()

    for index, row in df.iterrows():
        content  = row[content_col]
        time_val = row[time_col]
        if pd.isna(content):
            continue
        case_num = dl_extract_case_number(content)
        if not case_num:
            log(f"第{index+2}行：未提取到案号，跳过")
            continue
        time_str    = '' if pd.isna(time_val) else dl_sanitize(str(time_val).strip())
        folder_name = dl_sanitize(f"{case_num} {time_str}".strip())
        if folder_name in created:
            log(f"第{index+2}行：已处理过，跳过")
            continue
        folder_path = os.path.join(output_dir, folder_name)
        os.makedirs(folder_path, exist_ok=True)
        created.add(folder_name)
        log(f"\n第{index+2}行：{folder_name}")
        dl_process_sms(content, folder_path, log)

    log(f"\n✅ 全部完成，共处理 {len(created)} 条记录")
    log(f"📁 输出目录：{output_dir}")
    done()


# ══════════════════════════════════════════════
# 二、起诉状提取 - 文本读取模块
# ══════════════════════════════════════════════

def ex_ocr_pdf(path, log=print):
    if not HAS_PYMUPDF:
        log("    [OCR跳过] 请安装 PyMuPDF：pip install pymupdf")
        return ""
    if not HAS_EASYOCR:
        log("    [OCR跳过] 请安装 EasyOCR：pip install easyocr")
        return ""
    if not HAS_NUMPY:
        log("    [OCR跳过] 请安装 numpy：pip install numpy")
        return ""

    full_text = ""
    try:
        doc = fitz.open(path)
    except Exception as e:
        log(f"    [PDF打开失败] {e}")
        return ""

    reader = _get_ocr_reader(log)
    for page_num, page in enumerate(doc, start=1):
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
        if pix.n == 4:
            img = img[:, :, :3]
        try:
            results        = reader.readtext(img)
            results_sorted = sorted(results, key=lambda r: r[0][0][1])
            page_lines     = [item[1] for item in results_sorted if item[2] > 0.2]
            full_text     += "\n".join(page_lines) + "\n"
            log(f"    [OCR] 第 {page_num} 页完成，识别 {len(page_lines)} 行")
        except Exception as e:
            log(f"    [OCR第{page_num}页失败] {e}")

    doc.close()
    return full_text


def ex_read_pdf(path, log=print):
    if not HAS_PDF:
        return ""
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        log(f"    [PDF读取失败] {e}")

    if text.strip():
        return text

    log("    [提示] 检测到扫描版 PDF，启动 OCR 识别...")
    return ex_ocr_pdf(path, log)


def ex_read_docx(path):
    if not HAS_DOCX:
        return ""
    text = ""
    try:
        doc = python_docx.Document(path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
            text += "\n"
    except Exception:
        pass
    return text


def ex_read_txt(path):
    for enc in ("utf-8", "gbk", "gb2312", "utf-8-sig"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    return ""


def ex_extract_text(path, log=print):
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return ex_read_pdf(path, log)
    elif ext in (".docx", ".doc"):
        return ex_read_docx(path)
    elif ext == ".txt":
        return ex_read_txt(path)
    return ""


# ══════════════════════════════════════════════
# 三、起诉状提取 - 信息解析模块
# ══════════════════════════════════════════════

def ex_case_number_from_folder(folder_path):
    folder_name = os.path.basename(folder_path)
    m = re.search(r"[（(]\d{4}[）)]\S{2,30}?号", folder_name)
    return m.group().strip() if m else ""


def ex_format_court(raw):
    m = re.match(r"(\S+?)市(\S+?)人民法院", raw)
    if not m:
        return re.sub(r"人民法院$", "", raw).strip()
    city, district = m.group(1), m.group(2)
    if not any(district.endswith(s) for s in ("新区", "自治区", "自治县", "林区")):
        if district.endswith("区") or district.endswith("县"):
            district = district[:-1]
    return city + district


def ex_extract_court(folder_path):
    try:
        for filename in os.listdir(folder_path):
            if "短信内容" in filename and filename.lower().endswith(".txt"):
                content = ex_read_txt(os.path.join(folder_path, filename))
                m = re.search(r"[【\[](.+?)[】\]]", content)
                if m:
                    return ex_format_court(m.group(1).strip())
    except Exception:
        pass
    return ""


_PLAINTIFF_STOP_WORDS = (
    "身份证", "住所", "地址", "联系", "电话",
    "代理", "委托", "法定", "性别", "民族",
    "出生", "职业", "邮编", "原告在",
)


def _ex_clean_plaintiff_name(raw):
    raw = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])', '', raw)
    for kw in _PLAINTIFF_STOP_WORDS:
        idx = raw.find(kw)
        if idx != -1:
            raw = raw[:idx]
    return re.sub(r'[\s，,。、：:]+$', '', raw).strip()


def _ex_plaintiff_from_filename(filename):
    m = re.match(r'^([^诉]{2,15})诉', os.path.basename(filename))
    return m.group(1).strip() if m else ""


def ex_extract_plaintiffs(text, file_paths=None):
    names = []
    for m in re.finditer(
        r"原告[一二三四五六七八九十]?\s*[：:]\s*([^，,、\n]{2,25})", text
    ):
        name = _ex_clean_plaintiff_name(m.group(1))
        if len(name) >= 2 and name not in names:
            names.append(name)
    if not names and file_paths:
        for fp in file_paths:
            candidate = _ex_plaintiff_from_filename(fp)
            if candidate and candidate not in names:
                names.append(candidate)
    return names


def ex_extract_defendants(text):
    names = []
    for m in re.finditer(
        r"被告[一二三四五六七八九十]?\s*[：:]\s*([^，\n]{2,50})", text
    ):
        name = re.sub(r"[\s，,。]+$", "", m.group(1).strip())
        if name and name not in names:
            names.append(name)
    return names


def ex_extract_third_parties(text):
    names = []
    for m in re.finditer(
        r"第三人[一二三四五六七八九十]?\s*[：:]\s*([^，\n]{2,50})", text
    ):
        name = re.sub(r"[\s，,。]+$", "", m.group(1).strip())
        if name and name not in names:
            names.append(name)
    return names


def ex_clean(raw):
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


def ex_extract_claims(text):
    m = re.search(
        r"(?:诉讼请求|请求事项)\s*[：:]?\s*\n?(.*?)(?=事实(?:和|与|及)理由|此致|$)",
        text, re.DOTALL)
    return ex_clean(m.group(1)) if m else ""


def ex_extract_facts(text):
    m = re.search(
        r"事实(?:和|与|及)理由\s*[：:]?\s*\n?(.*?)(?=此致|具状人|附：|$)",
        text, re.DOTALL)
    return ex_clean(m.group(1)) if m else ""


# ══════════════════════════════════════════════
# 四、起诉状提取 - 被告排序 & 主体判断
# ══════════════════════════════════════════════

XINGYIN_KW = "行吟信息科技"

INTERNAL_COLS = ["案号","原告名字","其余被告","我方主体","被告1","一审法院",
                 "_空g_","诉讼请求","_空i_","_空j_","事实与理由"]
DISPLAY_COLS  = ["案号","原告名字","其余被告","我方主体","被告1","一审法院",
                 " ","诉讼请求","  ","   ","事实与理由"]
COL_WIDTHS    = {"案号":22,"原告名字":15,"其余被告":25,"我方主体":10,
                 "被告1":20,"一审法院":16," ":4,"诉讼请求":52,"  ":4,"   ":4,"事实与理由":62}


def ex_arrange_defendants(lst):
    found, others = False, []
    for d in lst:
        if XINGYIN_KW in d:
            found = True
        else:
            others.append(d)
    if found:
        return XINGYIN_KW, "、".join(others)
    return (lst[0], "、".join(lst[1:])) if lst else ("", "")


def ex_wo_fang(defendants, third_parties):
    return "第三人" if any(XINGYIN_KW in tp for tp in third_parties) else "被告"


# ══════════════════════════════════════════════
# 五、起诉状提取 - 文件夹重命名
# ══════════════════════════════════════════════

_CASE_START = re.compile(r"^[（(]\d{4}[）)]")


def ex_rename_folder(folder_path, plaintiff_name, log=print):
    folder_name = os.path.basename(folder_path)
    parent_dir  = os.path.dirname(folder_path)
    if not _CASE_START.match(folder_name):
        return folder_path
    prefix   = plaintiff_name if plaintiff_name else "【未识别】"
    new_path = os.path.join(parent_dir, prefix + folder_name)
    try:
        os.rename(folder_path, new_path)
        log(f"     [重命名] {folder_name} → {prefix + folder_name}")
        return new_path
    except FileExistsError:
        log(f"     [重命名跳过] 目标已存在：{prefix + folder_name}")
    except Exception as e:
        log(f"     [重命名失败] {e}")
    return folder_path


# ══════════════════════════════════════════════
# 六、起诉状提取 - 核心处理（按文件夹分组）
# ══════════════════════════════════════════════

SUPPORTED_EXT = {".pdf", ".docx", ".doc", ".txt"}


def ex_process_folder_group(folder_path, file_list, log):
    log(f"  📁 {os.path.basename(folder_path)}（{len(file_list)} 个文件）")
    combined_text = ""
    for fp in sorted(file_list):
        log(f"  -> {os.path.basename(fp)}")
        t = ex_extract_text(fp, log)
        if t.strip():
            combined_text += t + "\n"
            log(f"     [OK] 提取 {len(t.strip())} 字符")
        else:
            log("     [警告] 无可读文本")

    if not combined_text.strip():
        log("     [跳过] 所有文件均无可读文本")
        ex_rename_folder(folder_path, "", log)
        return None

    case_number   = ex_case_number_from_folder(folder_path)
    court         = ex_extract_court(folder_path)
    plaintiffs    = ex_extract_plaintiffs(combined_text, file_paths=file_list)
    defendants    = ex_extract_defendants(combined_text)
    third_parties = ex_extract_third_parties(combined_text)
    claims        = ex_extract_claims(combined_text)
    facts         = ex_extract_facts(combined_text)

    bei_gao_1, bei_gao_rest = ex_arrange_defendants(defendants)
    wo_fang       = ex_wo_fang(defendants, third_parties)
    plaintiff_str = "、".join(plaintiffs)

    log(f"     原告：{plaintiff_str or '(未识别)'}  法院：{court or '(未识别)'}")
    ex_rename_folder(folder_path, plaintiff_str, log)

    return {
        "案号": case_number, "原告名字": plaintiff_str,
        "其余被告": bei_gao_rest, "我方主体": wo_fang,
        "被告1": bei_gao_1, "一审法院": court,
        "_空g_": "", "诉讼请求": claims,
        "_空i_": "", "_空j_": "", "事实与理由": facts,
    }


def ex_scan_folder(root_path, log):
    folder_files = defaultdict(list)
    for root, _, files in os.walk(root_path):
        for filename in sorted(files):
            if "起诉状" not in filename:
                continue
            if "短信内容" in filename:
                continue
            if Path(filename).suffix.lower() not in SUPPORTED_EXT:
                continue
            folder_files[root].append(os.path.join(root, filename))

    total_files   = sum(len(v) for v in folder_files.values())
    total_folders = len(folder_files)
    log(f"共找到 {total_files} 个起诉状文件，分布于 {total_folders} 个文件夹\n")

    records = []
    for folder_path, file_list in folder_files.items():
        result = ex_process_folder_group(folder_path, file_list, log)
        if result:
            records.append(result)
        log("")
    return records


# ══════════════════════════════════════════════
# 七、起诉状提取 - Excel 写入
# ══════════════════════════════════════════════

def _is_file_locked(filepath):
    if not os.path.exists(filepath):
        return False
    try:
        with open(filepath, "a"):
            return False
    except PermissionError:
        return True


def _get_available_path(base_path):
    if not _is_file_locked(base_path):
        return base_path
    p = Path(base_path)
    for i in range(1, 100):
        candidate = p.parent / f"{p.stem}_{i}{p.suffix}"
        if not _is_file_locked(str(candidate)):
            return str(candidate)
    raise RuntimeError("无法找到可用输出路径，请关闭 Excel 后重试。")


def _ex_set_style(writer, df_display):
    from openpyxl.styles import Font, Alignment, PatternFill
    ws = writer.sheets["起诉状汇总"]
    for idx, col_name in enumerate(df_display.columns, start=1):
        col_letter = ws.cell(row=1, column=idx).column_letter
        ws.column_dimensions[col_letter].width = COL_WIDTHS.get(col_name, 15)
    header_fill = PatternFill(fill_type="solid", fgColor="DDEEFF")
    for cell in ws[1]:
        cell.font      = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.fill      = header_fill
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "A2"


# ══════════════════════════════════════════════
# 八、起诉状提取 - 任务入口
# ══════════════════════════════════════════════

def ex_run_task(folder_path, log, done):
    records = ex_scan_folder(folder_path, log)
    if not records:
        log('未找到任何含"起诉状"的可处理文件')
        done()
        return

    df         = pd.DataFrame(records, columns=INTERNAL_COLS)
    df_display = df.copy()
    df_display.columns = DISPLAY_COLS

    base_output = os.path.join(folder_path, "起诉状信息汇总.xlsx")
    try:
        actual_path = _get_available_path(base_output)
        if actual_path != base_output:
            log(f"⚠️  原文件被占用，另存为：{os.path.basename(actual_path)}")
        with pd.ExcelWriter(actual_path, engine="openpyxl") as writer:
            df_display.to_excel(writer, index=False, sheet_name="起诉状汇总")
            _ex_set_style(writer, df_display)
        log(f"✅ 完成！共处理 {len(records)} 个文件夹")
        log(f"📁 输出：{actual_path}")
    except Exception as e:
        log(f"❌ 保存Excel失败：{e}")
    done()


# ══════════════════════════════════════════════
# 新增：初筛-判决裁定书 - 业务逻辑
# ══════════════════════════════════════════════

_SC_PREFIXES = ("【传票】", "【裁定书】", "【判决书】", "【调解书】")   # ← 新增【调解书】


def sc_detect_pdf_type(folder_path: str) -> Tuple[bool, bool, bool, bool]:
    """
    扫描文件夹内 PDF 文件名，返回 (has_caiding, has_panjue, has_tiaojie, has_chuanpiao)。
    """
    has_caiding   = False
    has_panjue    = False
    has_tiaojie   = False
    has_chuanpiao = False          # ← 新增
    try:
        for filename in os.listdir(folder_path):
            if not filename.lower().endswith(".pdf"):
                continue
            if "裁定书" in filename or "本院法律文书正本_1" in filename:
                has_caiding = True
            if "判决书" in filename:
                has_panjue = True
            if "调解书" in filename:
                has_tiaojie = True
            if "传票" in filename:             # ← 新增
                has_chuanpiao = True
    except PermissionError:
        pass
    return has_caiding, has_panjue, has_tiaojie, has_chuanpiao   # ← 新增返回值



def sc_determine_prefix(
    has_caiding: bool,
    has_panjue: bool,
    has_tiaojie: bool,
    has_chuanpiao: bool,           # ← 新增参数
) -> Optional[str]:
    """
    优先级：传票 > 裁定书 > 判决书 > 调解书
    """
    if has_chuanpiao:              # ← 最高优先级
        return "【传票】"
    if has_caiding:
        return "【裁定书】"
    if has_panjue:
        return "【判决书】"
    if has_tiaojie:
        return "【调解书】"
    return None


def sc_run_task(target_folder: str, log, done):
    """初筛任务主函数，供后台线程调用。"""
    if not os.path.isdir(target_folder):
        log(f"❌ 目标文件夹不存在: {target_folder}")
        done()
        return

    log("=" * 52)
    log(f"目标文件夹: {target_folder}")
    log("=" * 52)

    stats = {
        "renamed":  0,
        "skipped":  0,
        "no_match": 0,
        "conflict": 0,
        "error":    0,
    }

    try:
        subfolders = [
            item for item in os.listdir(target_folder)
            if os.path.isdir(os.path.join(target_folder, item))
        ]
    except Exception as e:
        log(f"❌ 无法读取目标文件夹: {e}")
        done()
        return

    log(f"共发现子文件夹: {len(subfolders)} 个\n")

    for folder_name in subfolders:
        folder_path = os.path.join(target_folder, folder_name)
        log(f"处理: {folder_name}")

        # 已有任意前缀则跳过
        if any(folder_name.startswith(p) for p in _SC_PREFIXES):
            log("  → 已有前缀，跳过\n")
            stats["skipped"] += 1
            continue

        try:
            # sc_run_task 内，try 块中完整替换为：
            has_caiding, has_panjue, has_tiaojie, has_chuanpiao = sc_detect_pdf_type(folder_path)

            types_found = []
            if has_chuanpiao:
                types_found.append("传票")
            if has_caiding:
                types_found.append("裁定书")
            if has_panjue:
                types_found.append("判决书")
            if has_tiaojie:
                types_found.append("调解书")

            if len(types_found) > 1:
                top = types_found[0]  # 列表顺序即优先级顺序
                log(f"  ⚠ 同时含有 {'、'.join(types_found)} PDF，按优先级取【{top}】前缀")

            prefix = sc_determine_prefix(has_caiding, has_panjue, has_tiaojie, has_chuanpiao)

            if prefix is None:
                log("  → 未找到裁定书、判决书、调解书或传票 PDF，跳过\n")
                stats["no_match"] += 1
                continue

            new_name = prefix + folder_name
            new_path = os.path.join(target_folder, new_name)

            if os.path.exists(new_path):
                log(f"  → ⚠ 目标路径已存在，跳过: {new_name}\n")
                stats["conflict"] += 1
                continue

            os.rename(folder_path, new_path)
            log(f"  → ✓ 重命名成功:")
            log(f"       原名: {folder_name}")
            log(f"       新名: {new_name}\n")
            stats["renamed"] += 1

        except Exception as e:
            log(f"  → ✗ 处理出错: {e}\n")
            stats["error"] += 1

    log("=" * 52)
    log("处理完成！汇总报告：")
    log(f"  ✓ 成功重命名   : {stats['renamed']}  个")
    log(f"  - 已有前缀跳过 : {stats['skipped']}  个")
    log(f"  - 无匹配PDF跳过: {stats['no_match']} 个")
    log(f"  ⚠ 目标路径冲突 : {stats['conflict']} 个")
    log(f"  ✗ 处理出错     : {stats['error']}  个")
    log("=" * 52)
    done()

# ══════════════════════════════════════════════
# 九、通用日志组件
# ══════════════════════════════════════════════

def make_log_box(parent):
    frame = tk.Frame(parent)
    text  = tk.Text(frame, font=("Consolas", 9),
                    bg="#1e1e1e", fg="#d4d4d4",
                    wrap="word", state="disabled")
    sb    = tk.Scrollbar(frame, command=text.yview)
    text.configure(yscrollcommand=sb.set)
    text.pack(side="left", fill="both", expand=True)
    sb.pack(side="right", fill="y")

    def log(msg):
        def _w():
            text.configure(state="normal")
            text.insert("end", str(msg) + "\n")
            text.see("end")
            text.configure(state="disabled")
        try:
            text.after(0, _w)
        except Exception:
            pass

    def clear():
        text.configure(state="normal")
        text.delete("1.0", "end")
        text.configure(state="disabled")

    return frame, log, clear


# ══════════════════════════════════════════════
# 十、Tab1 - 文书下载界面
# ══════════════════════════════════════════════

class DownloadTab(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self._build()

    def _build(self):
        param = tk.LabelFrame(self, text="参数设置",
                              font=("微软雅黑", 10), padx=10, pady=8)
        param.pack(fill="x", padx=15, pady=10)

        tk.Label(param, text="Excel 文件：", width=12, anchor="e").grid(row=0, column=0, pady=4, sticky="e")
        self.excel_var = tk.StringVar()
        tk.Entry(param, textvariable=self.excel_var, width=44).grid(row=0, column=1, padx=5, sticky="ew")
        tk.Button(param, text="浏览", width=6, command=self._pick_excel).grid(row=0, column=2)

        tk.Label(param, text="内容列名：", width=12, anchor="e").grid(row=1, column=0, pady=4, sticky="e")
        self.content_col = tk.StringVar(value="Message Body")
        tk.Entry(param, textvariable=self.content_col, width=44).grid(row=1, column=1, padx=5, sticky="ew")

        tk.Label(param, text="时间列名：", width=12, anchor="e").grid(row=2, column=0, pady=4, sticky="e")
        self.time_col = tk.StringVar(value="Date")
        tk.Entry(param, textvariable=self.time_col, width=44).grid(row=2, column=1, padx=5, sticky="ew")

        tk.Label(param, text="输出目录：", width=12, anchor="e").grid(row=3, column=0, pady=4, sticky="e")
        self.output_var = tk.StringVar()
        tk.Entry(param, textvariable=self.output_var, width=44).grid(row=3, column=1, padx=5, sticky="ew")
        tk.Button(param, text="浏览", width=6, command=self._pick_output).grid(row=3, column=2)
        param.columnconfigure(1, weight=1)

        tk.Label(self, text="💡 列名请参照Excel第一行表头填写",
                 fg="gray", font=("微软雅黑", 9)).pack(anchor="w", padx=20)

        btn_f = tk.Frame(self)
        btn_f.pack(pady=6)
        self.run_btn = tk.Button(btn_f, text="▶  开始执行", width=14, height=2,
                                 bg="#2c7be5", fg="white",
                                 font=("微软雅黑", 11, "bold"),
                                 command=self._start)
        self.run_btn.pack(side="left", padx=5)
        tk.Button(btn_f, text="清空日志", width=10, height=2,
                  command=lambda: self._clear()).pack(side="left", padx=5)

        self.progress = ttk.Progressbar(self, mode="indeterminate")
        self.progress.pack(fill="x", padx=15)

        log_frame = tk.LabelFrame(self, text="运行日志",
                                  font=("微软雅黑", 10), padx=5, pady=5)
        log_frame.pack(fill="both", expand=True, padx=15, pady=8)
        box, self._log, self._clear = make_log_box(log_frame)
        box.pack(fill="both", expand=True)

    def _pick_excel(self):
        p = filedialog.askopenfilename(
            title="选择 Excel",
            filetypes=[("Excel", "*.xlsx *.xls"), ("所有", "*.*")])
        if p:
            self.excel_var.set(p)
            if not self.output_var.get():
                self.output_var.set(os.path.dirname(p) + "\\输出文件")

    def _pick_output(self):
        p = filedialog.askdirectory(title="选择输出目录")
        if p:
            self.output_var.set(p)

    def _start(self):
        excel  = self.excel_var.get().strip()
        c_col  = self.content_col.get().strip()
        t_col  = self.time_col.get().strip()
        outdir = self.output_var.get().strip()

        if not excel:
            messagebox.showwarning("提示", "请选择 Excel 文件")
            return
        if not os.path.exists(excel):
            messagebox.showerror("错误", f"文件不存在：{excel}")
            return
        if not outdir:
            messagebox.showwarning("提示", "请选择输出目录")
            return

        self.run_btn.configure(state="disabled", text="执行中...")
        self.progress.start(10)
        self._log(f"▶ 开始执行\n  Excel：{excel}\n  输出：{outdir}\n")

        threading.Thread(
            target=dl_run_task,
            args=(excel, c_col, t_col, outdir, self._log, self._done),
            daemon=True
        ).start()

    def _done(self):
        self.after(0, lambda: (
            self.progress.stop(),
            self.run_btn.configure(state="normal", text="▶  开始执行")
        ))


# ══════════════════════════════════════════════
# 十一、Tab2 - 起诉状提取界面
# ══════════════════════════════════════════════

class ExtractTab(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self._build()

    def _build(self):
        param = tk.LabelFrame(self, text="参数设置",
                              font=("微软雅黑", 10), padx=10, pady=8)
        param.pack(fill="x", padx=15, pady=10)

        tk.Label(param, text="案件文件夹：", width=12, anchor="e").grid(row=0, column=0, pady=4, sticky="e")
        self.folder_var = tk.StringVar()
        tk.Entry(param, textvariable=self.folder_var, width=44).grid(row=0, column=1, padx=5, sticky="ew")
        tk.Button(param, text="浏览", width=6, command=self._pick_folder).grid(row=0, column=2)
        param.columnconfigure(1, weight=1)

        tk.Label(self,
                 text='💡 选择包含多个案件子文件夹的根目录，程序自动扫描所有"起诉状"文件',
                 fg="gray", font=("微软雅黑", 9), wraplength=580, justify="left"
                 ).pack(anchor="w", padx=20)

        dep_f = tk.Frame(self)
        dep_f.pack(anchor="w", padx=20, pady=2)

        def dep_label(parent, ok, name, pip_name):
            color = "#2e7d32" if ok else "#c62828"
            text  = f"{'✓' if ok else '✗'} {name}{'  ' if ok else f'  pip install {pip_name}'}"
            tk.Label(parent, text=text, fg=color, font=("微软雅黑", 9)).pack(side="left", padx=6)

        dep_label(dep_f, HAS_PDF,     "pdfplumber",  "pdfplumber")
        dep_label(dep_f, HAS_DOCX,    "python-docx", "python-docx")
        dep_label(dep_f, HAS_PYMUPDF, "PyMuPDF",     "pymupdf")
        dep_label(dep_f, HAS_EASYOCR, "EasyOCR",     "easyocr")
        dep_label(dep_f, HAS_NUMPY,   "numpy",        "numpy")

        ocr_tip = ("（已具备扫描版 PDF OCR 能力）"
                   if (HAS_PYMUPDF and HAS_EASYOCR and HAS_NUMPY)
                   else "（安装 PyMuPDF + EasyOCR + numpy 可支持扫描版 PDF）")
        tk.Label(self, text=ocr_tip, fg="gray", font=("微软雅黑", 8)
                 ).pack(anchor="w", padx=20)

        btn_f = tk.Frame(self)
        btn_f.pack(pady=6)
        self.run_btn = tk.Button(btn_f, text="▶  开始提取", width=14, height=2,
                                 bg="#2e7d32", fg="white",
                                 font=("微软雅黑", 11, "bold"),
                                 command=self._start)
        self.run_btn.pack(side="left", padx=5)
        tk.Button(btn_f, text="清空日志", width=10, height=2,
                  command=lambda: self._clear()).pack(side="left", padx=5)

        self.progress = ttk.Progressbar(self, mode="indeterminate")
        self.progress.pack(fill="x", padx=15)

        log_frame = tk.LabelFrame(self, text="运行日志",
                                  font=("微软雅黑", 10), padx=5, pady=5)
        log_frame.pack(fill="both", expand=True, padx=15, pady=8)
        box, self._log, self._clear = make_log_box(log_frame)
        box.pack(fill="both", expand=True)

    def _pick_folder(self):
        p = filedialog.askdirectory(title="选择案件根目录")
        if p:
            self.folder_var.set(p)

    def _start(self):
        folder = self.folder_var.get().strip()
        if not folder:
            messagebox.showwarning("提示", "请选择案件文件夹")
            return
        if not os.path.isdir(folder):
            messagebox.showerror("错误", f"目录不存在：{folder}")
            return

        self.run_btn.configure(state="disabled", text="提取中...")
        self.progress.start(10)
        self._log(f"▶ 开始扫描\n  目录：{folder}\n")

        threading.Thread(
            target=ex_run_task,
            args=(folder, self._log, self._done),
            daemon=True
        ).start()

    def _done(self):
        self.after(0, lambda: (
            self.progress.stop(),
            self.run_btn.configure(state="normal", text="▶  开始提取")
        ))


# ══════════════════════════════════════════════
# 十二、Tab3 - 初筛判决/裁定书界面
# ══════════════════════════════════════════════

class ScreeningTab(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self._build()

    def _build(self):
        # ── 参数区 ──────────────────────────────
        param = tk.LabelFrame(self, text="参数设置",
                              font=("微软雅黑", 10), padx=10, pady=8)
        param.pack(fill="x", padx=15, pady=10)

        tk.Label(param, text="案件根目录：", width=12, anchor="e").grid(
            row=0, column=0, pady=4, sticky="e")
        self.folder_var = tk.StringVar()
        tk.Entry(param, textvariable=self.folder_var, width=44).grid(
            row=0, column=1, padx=5, sticky="ew")
        tk.Button(param, text="浏览", width=6,
                  command=self._pick_folder).grid(row=0, column=2)
        param.columnconfigure(1, weight=1)

        # ── 说明文字（逐行写，避免括号内多行字符串触发低版本误报）──
        tip_lines = (
            "💡 扫描所选目录下的直接子文件夹，根据子文件夹内 PDF 文件名添加前缀：",
            "   · 含「传票」PDF      →  【传票】 （最高优先级）",
            "   · 含「裁定书」PDF    →  【裁定书】",
            "   · 含「判决书」PDF    →  【判决书】",
            "   · 含「调解书」PDF    →  【调解书】",
            "   · 同时含多种时按上述优先级取最高；已有前缀的文件夹自动跳过",
        )
        for line in tip_lines:
            tk.Label(self, text=line, fg="gray",
                     font=("微软雅黑", 9), anchor="w").pack(
                anchor="w", padx=20)

        # ── 按钮区 ──────────────────────────────
        btn_f = tk.Frame(self)
        btn_f.pack(pady=6)
        self.run_btn = tk.Button(
            btn_f, text="▶  开始初筛", width=14, height=2,
            bg="#6a1b9a", fg="white",
            font=("微软雅黑", 11, "bold"),
            command=self._start,
        )
        self.run_btn.pack(side="left", padx=5)
        tk.Button(btn_f, text="清空日志", width=10, height=2,
                  command=lambda: self._clear()).pack(side="left", padx=5)

        # ── 进度条 ──────────────────────────────
        self.progress = ttk.Progressbar(self, mode="indeterminate")
        self.progress.pack(fill="x", padx=15)

        # ── 日志区 ──────────────────────────────
        log_frame = tk.LabelFrame(self, text="运行日志",
                                  font=("微软雅黑", 10), padx=5, pady=5)
        log_frame.pack(fill="both", expand=True, padx=15, pady=8)
        box, self._log, self._clear = make_log_box(log_frame)
        box.pack(fill="both", expand=True)

    def _pick_folder(self):
        p = filedialog.askdirectory(title="选择案件根目录")
        if p:
            self.folder_var.set(p)

    def _start(self):
        folder = self.folder_var.get().strip()
        if not folder:
            messagebox.showwarning("提示", "请选择案件根目录")
            return
        if not os.path.isdir(folder):
            messagebox.showerror("错误", f"目录不存在：{folder}")
            return

        self.run_btn.configure(state="disabled", text="初筛中...")
        self.progress.start(10)
        self._log(f"▶ 开始初筛\n  目录：{folder}\n")

        threading.Thread(
            target=sc_run_task,
            args=(folder, self._log, self._done),
            daemon=True,
        ).start()

    def _done(self):
        self.after(0, lambda: (
            self.progress.stop(),
            self.run_btn.configure(state="normal", text="▶  开始初筛"),
        ))


# ══════════════════════════════════════════════
# 十三、主窗口
# ══════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("法院工具箱")
        self.geometry("700x680")
        self.minsize(620, 540)
        self._build()

    def _build(self):
        tk.Label(self, text="⚖  法院工具箱",
                 font=("微软雅黑", 17, "bold"),
                 fg="#1a3c5e").pack(pady=(12, 0))
        tk.Label(self, text="Court Document Toolkit",
                 font=("微软雅黑", 9), fg="gray").pack(pady=(0, 8))

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        tab1 = DownloadTab(nb)
        tab2 = ExtractTab(nb)
        tab3 = ScreeningTab(nb)

        nb.add(tab1, text="  📥  文书下载  ")
        nb.add(tab2, text="  📄  起诉状提取  ")
        nb.add(tab3, text="  🔍  初筛裁判书  ")


if __name__ == "__main__":
    App().mainloop()
